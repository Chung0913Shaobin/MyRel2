from flask import Flask, request, render_template_string, send_file, session
import requests
import os
import logging
from docx import Document
from PyPDF2 import PdfReader
from werkzeug.utils import secure_filename
import google.generativeai as genai
import json
import re
from fpdf import FPDF


app = Flask(__name__)
app.secret_key = "this_is_a_test_key_for_dev_123" 
OUTPUT_FOLDER = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

DEEPSEEK_API_KEY = "sk-c17852df776c481eaf9c7d7e7951b498"
GEMINI_API_KEY = "AIzaSyB2AsIs5YNb5wF_wpZ2xVySwv3mZ5pRF5s"
OPENAI_API_KEY = "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"

DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
OPENAI_API_URL = "https://api.openai.com/v1/chat/completions"
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1/models/gemini-pro:generateContent"



genai.configure(api_key=GEMINI_API_KEY)
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')

def extract_text(file_path, extension):
    try:
        if extension == 'pdf':
            reader = PdfReader(file_path)
            return "".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif extension == 'docx':
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            return ""
    except Exception as e:
        logging.error(f"提取檔案內容失敗：{e}")
        return ""

def ask_deepseek(prompt):
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 2048
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload)
        if response.status_code != 200:
            logging.error(f"DeepSeek API 回應錯誤：{response.status_code} - {response.text}")
            return "DeepSeek API 回應錯誤，請稍後再試。"

        data = response.json()

        if "choices" not in data or not data["choices"]:
            logging.error(f"DeepSeek 回傳格式錯誤：{data}")
            return "DeepSeek API 回傳內容異常。"

        return data["choices"][0]["message"]["content"].strip()

    except Exception as e:
        logging.error(f"與 DeepSeek 通訊時發生錯誤：{e}")
        return "與 DeepSeek 通訊失敗，請稍後再試。"

def ask_openai(prompt):
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "gpt-4",
        "messages": [{"role": "user", "content": prompt}]
    }

    try:
        response = requests.post(OPENAI_API_URL, headers=headers, json=payload)
        if response.status_code != 200:
            logging.error(f"OpenAI API 回應錯誤：{response.status_code} - {response.text}")
            return "OpenAI API 回應錯誤，請稍後再試。"

        data = response.json()

        if "choices" not in data or not data["choices"]:
            logging.error(f"OpenAI API 回傳格式錯誤：{data}")
            return "OpenAI API 回傳內容異常。"

        return data["choices"][0]["message"]["content"].strip()
    
    except Exception as e:
        logging.error(f"與 OpenAI 通訊時發生錯誤：{e}")
        return "與 OpenAI 通訊失敗，請稍後再試。"

def ask_gemini(prompt):
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)

        if not response.text:
            logging.error(f"Gemini 回傳內容為空或異常：{response}")
            return "Gemini API 回傳內容異常。"

        return response.text.strip()

    except Exception as e:
        logging.error(f"與 Gemini 通訊時發生錯誤：{e}")
        return "與 Gemini 通訊失敗，請稍後再試。"

def clean_ai_output(text):
    import re
    text = re.sub(r"#+\s*", "", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  
    text = re.sub(r"-{2,}", "", text)
    text = re.sub(r"^[-•]\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()

def save_analysis_to_word(text, filepath=None):
    if filepath is None:
        filepath = os.path.join(OUTPUT_FOLDER, "result.docx")
    doc = Document()
    doc.add_heading("履歷分析結果", 0)
    for line in text.strip().split("\n"):
        doc.add_paragraph(line)
    doc.save(filepath)
    print(f"Word 儲存成功：{filepath}")
    print("建立 Word 檔案於：", filepath)
    return filepath
    

def save_analysis_to_pdf(text, filepath=None):
    if filepath is None:
        filepath = os.path.join(OUTPUT_FOLDER, "result.pdf")

    pdf = FPDF()
    pdf.add_page()

    font_path = "C:/Windows/Fonts/msjh.ttc"
    pdf.add_font("MSJH", "", font_path, uni=True)
    pdf.set_font("MSJH", "", 14)

    pdf.multi_cell(0, 10, "履歷分析結果", align="C")
    pdf.ln(10)

    for line in text.strip().split("\n"):
        pdf.multi_cell(0, 10, line)

    pdf.output(filepath)
    print(f"PDF 儲存成功：{filepath}")
    return filepath




@app.route("/", methods=["GET"])
def index():
    return '''
    <!DOCTYPE html>
    <html lang="zh-TW">
    <head>
        <meta charset="UTF-8">
        <title>上傳您的履歷</title>
        <style>
            body {
                font-family: 'Noto Sans TC', sans-serif;
                background-color: #f5f5f7;
                padding: 50px;
                text-align: center;
            }
            .upload-box {
                border: 2px dashed #1a4b8c;
                padding: 180px;
                border-radius: 12px;
                background-color: #fff;
                display: flex;
                flex-direction: column;
                align-items: center;
                justify-content: center;
                margin: 20px auto;
                width: 70%;
                cursor: pointer;
            }
            .upload-box p {
                margin: 10px 0;
                font-size: 1.4em;
                color: #1a4b8c;
            }
            .upload-box p:last-child {
                font-size: 1.2em;
                color: #7a7a7a;
            }
            input[type="file"] {
                display: none;
            }
            input[type="submit"] {
                background-color: #1a4b8c;
                color: white;
                border: none;
                padding: 12px 30px;
                font-size: 1.4em;
                border-radius: 8px;
                cursor: pointer;
                margin-top: 20px;
            }
            #loading {
                display: none;
                margin-top: 20px;
                font-size: 1.4em;
                color: #555;
            }
        </style>
        <script>
    function showLoading() {
        document.getElementById('loading').style.display = 'block';
    }
    function triggerFileUpload() {
        document.getElementById('file-input').click();
    }
    document.addEventListener('DOMContentLoaded', () => {
        const fileInput = document.getElementById('file-input');
        const uploadBoxText = document.querySelector('.upload-box p');
        fileInput.addEventListener('change', () => {
            if (fileInput.files.length > 0) {
                uploadBoxText.textContent = `${fileInput.files[0].name} (${(fileInput.files[0].size / 1024).toFixed(2)} KB)`;
            }
        });
    });
</script>
    </head>
    <body>
        <h1>上傳您的履歷</h1>
        <form method="post" action="/upload" enctype="multipart/form-data" onsubmit="showLoading()">
            <div class="upload-box" onclick="triggerFileUpload()">
                <svg xmlns="http://www.w3.org/2000/svg" width="48" height="48" fill="#1a4b8c" viewBox="0 0 24 24"><path d="M12 0l3.09 6.26L22 7.27l-5 4.87L18.18 18 12 15.27 5.82 18 7 12.14l-5-4.87 6.91-1.01L12 0z"/></svg>
                <p>點擊或拖放檔案至此處上傳</p>
                <p>支援 PDF、Word 檔案格式</p>
            </div>
            <input id="file-input" type="file" name="file" accept=".pdf,.docx">
            <br>
            <input type="submit" value="上傳並分析">
            <div id="loading">
                🔄 分析中，請稍候...
                <div style="width: 100%; background-color: #ddd; border-radius: 8px; margin-top: 10px;">
                    <div style="width: 100%; height: 20px; background-color: #1a4b8c; border-radius: 8px; animation: progress 3s infinite;"></div>
                </div>
            </div>
        </form>

        <style>
            @keyframes progress {
                0% { width: 0%; }
                50% { width: 50%; }
                100% { width: 100%; }
            }
        </style>
    </body>
    </html>
    '''

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "未選擇檔案"
    file = request.files["file"]
    if file.filename == "":
        return "未選擇檔案"
    filename = secure_filename(file.filename)
    file_ext = filename.rsplit(".", 1)[1].lower()
    if file_ext not in ["pdf", "docx"]:
        return "不支援的檔案格式"
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)
    resume_text = extract_text(filepath, file_ext)
    if not resume_text.strip():
        return "履歷內容為空或無法解析"

    skill_prompt = f"""
    以下是我的履歷內容。請提取我具備的技能，並根據技能的特性進行分類（如編程語言、框架與工具、數據分析等）。
    請忽略文中提到的不熟悉或不擅長的技能，只列出明確提到具備或熟練的技能。
    輸出格式如下：

    分類名稱1：技能1, 技能2, 技能3

    分類名稱2：技能4, 技能5

    分類名稱3：技能6, 技能7
    ...

    履歷內容：
    {resume_text}
    """

    suggestion_prompt = f"""
    請根據以下履歷內容，分析並提供提升履歷質量的具體建議。

    履歷內容：
    {resume_text}

    請針對以下方面提出建議：
    - 段落結構
    - 資訊完整性
    - 技能表達
    - 專業用詞
    - 排版設計
    - 個人形象塑造

    請將建議整理為三個重點，每點建議控制在 50 字內，簡明扼要地提供具體改進方向。

    履歷優化建議：
    1. [建議 1]
    2. [建議 2]
    3. [建議 3]
    """
    # print("DEBUG - 履歷前 500 字：\n", resume_text[:500])

    evaluation_prompt = f"""
    你是一位專業的人力資源顧問，請根據以下履歷進行分析與評分，並根據履歷內容推薦適合的職缺。

    請嚴格依據履歷中**實際出現的技能、經歷、專案內容**，判斷其完整性與符合度。
    ❗ 請避免根據既有模板或平均分配給分數，務必根據具體內容差異化評估。

    請依以下步驟執行分析：

    一、依照以下評分項目給出分數，滿分為 25 分，每項需根據摘要細節給分
    二、請所有分數皆以整數表示，若有小數點則四捨五入
    三、回傳格式為標準 JSON，如下：

    ```json
    {{
    "技能匹配度": 16,
    "工作經驗與成就": 17,
    "履歷結構與可讀性": 23,
    "個人特色": 20,
    "總分": 76,
    }}

    請依據評分提供以下資訊：

    履歷總分：
    - 技能匹配度: [分數]
    - 工作經驗與成就: [分數]
    - 履歷結構與可讀性: [分數]
    - 個人特色: [分數]
    - 總分: [總分數]

    適合的職缺：
    1. [職缺名稱]
    2. [職缺名稱]
    3. [職缺名稱]

    關鍵字與技能優化建議：
    請根據推薦的職缺，提供可補充或優化的技能與關鍵字，以提升履歷的市場競爭力。

    - 建議補充技能：
    - [技能名稱]: [該技能對應的職缺優勢]
    - [技能名稱]: [該技能對應的職缺優勢]

    - 建議加入關鍵字：
    - [關鍵字]: [關鍵字對履歷吸引力的提升效果]
    - [關鍵字]: [關鍵字對履歷吸引力的提升效果]

    履歷內容：
    {resume_text}
    """

    skills_d = ask_deepseek(skill_prompt)
    suggestions_d = ask_deepseek(suggestion_prompt)
    evaluation_d = ask_deepseek(evaluation_prompt)

    skills_g = ask_gemini(skill_prompt)
    suggestions_g = ask_gemini(suggestion_prompt)
    evaluation_g = ask_gemini(evaluation_prompt)

    skills_o = ask_openai(skill_prompt)
    suggestions_o = ask_openai(suggestion_prompt)
    evaluation_o = ask_openai(evaluation_prompt)

    def extract_json_from_text(text):

        try:
            # 嘗試抓出 ```json 區塊
            json_block = re.search(r'```json\s*({[\s\S]+?})\s*```', text)
            raw = json_block.group(1) if json_block else text

            # 嘗試修復 JSON：
            raw = raw.replace("：", ":")  # 中文冒號轉英文冒號
            raw = re.sub(r"([{\[,])\s*([\u4e00-\u9fa5\w]+)\s*:", r'\1 "\2":', raw)  # 補雙引號到 key
            raw = re.sub(r":\s*'([^']*)'", r': "\1"', raw)  # 單引號 value 改雙引號
            raw = re.sub(r"//.*", "", raw)  # 移除註解
            raw = re.sub(r",\s*}", "}", raw)  # 移除結尾多餘逗號
            raw = re.sub(r",\s*]", "]", raw)

            # 嘗試讀取成 JSON
            return json.loads(raw)

        except Exception as e:
            logging.error(f"從文字中擷取 JSON 失敗：{e}")
            return {
                "技能匹配度": 0,
                "工作經驗與成就": 0,
                "履歷結構與可讀性": 0,
                "個人特色": 0,
                "總分": 0
            }
        
    def average_scores(*score_dicts):
        keys = ["技能匹配度", "工作經驗與成就", "履歷結構與可讀性", "個人特色"]
        valid_dicts = [sd for sd in score_dicts if sd.get("總分", 0) >= 10]

        if not valid_dicts:
            return {k: 0 for k in keys + ["總分"]}

        avg = {}
        for key in keys:
            total = sum(sd.get(key, 0) for sd in valid_dicts)
            avg[key] = round(total / len(valid_dicts))
        avg["總分"] = sum(avg.values())
        return avg
    
    final_prompt = f"""以下是三個 AI 模型對履歷的分析結果，請統整成最佳建議：

    技能分類：
    - DeepSeek: {skills_d}
    - Gemini: {skills_g}
    - OpenAI: {skills_o}

    履歷建議：
    - DeepSeek: {suggestions_d}
    - Gemini: {suggestions_g}
    - OpenAI: {suggestions_o}

    履歷評分與推薦職缺：
    - DeepSeek: {evaluation_d}
    - Gemini: {evaluation_g}
    - OpenAI: {evaluation_o}

    請統整上述資訊，合併重點並避免重複，提供一份清楚、具體且精煉的分析與建議。"""
    
    
    final_result = ask_deepseek(final_prompt)
    cleaned_result = clean_ai_output(final_result)

    score_data_d = extract_json_from_text(evaluation_d)
    score_data_g = extract_json_from_text(evaluation_g)
    score_data_o = extract_json_from_text(evaluation_o)

    score_data = average_scores(score_data_d, score_data_g, score_data_o)

    with open(os.path.join(OUTPUT_FOLDER, "result.txt"), "w", encoding="utf-8") as f:
        f.write(cleaned_result)

    return render_template_string(f'''
    <!DOCTYPE html>
    <html lang="zh-TW">
    <head>
        <meta charset="UTF-8">
        <title>履歷分析結果</title>
        <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
        <style>
            body {{
                font-family: 'Noto Sans TC', sans-serif;
                background-color: #f5f5f7;
                padding: 30px;
                position: relative;
            }}
            .h1 {{
                color: #1a4b8c;
                text-align: center;
                margin-bottom: 40px;
                font-size: 2.6em;
            }}
            .section {{
                background: white;
                padding: 30px;
                border-radius: 12px;
                margin-bottom: 30px;
            }}
            .with-chart-line {{
                display: flex;
                gap: 30px;
                align-items: flex-start;
            }}
            .text-content {{
                flex: 1;
                font-size: 1.2em;
                line-height: 1.6;
                white-space: pre-wrap;
                word-break: break-word;
            }}
            .right-charts {{
                display: flex;
                flex-direction: column;
                gap: 30px;
                width: 400px;
                min-width: 300px;
            }}
            .chart-container {{
                background-color: white;
                padding: 20px;
                border-radius: 12px;
                box-shadow: 0 0 10px rgba(0,0,0,0.05);
                text-align: center;
            }}
            .chart-title {{
                color: #1a4b8c;
                font-weight: bold;
                margin-bottom: 10px;
            }}
            /* 加入右上角下載按鈕區塊 */
            .download-buttons {{
                position: absolute;
                top: 20px;
                right: 40px;
                z-index: 999;
            }}
            .download-buttons a {{
                display: inline-block;
                padding: 10px 16px;
                margin-left: 10px;
                border-radius: 8px;
                text-decoration: none;
                font-weight: bold;
                color: white;
            }}
            .btn-word {{
                background-color: #1a4b8c;
            }}
            .btn-pdf {{
                background-color: #1a4b8c;
            }}
        </style>
    </head>
    <body>

        <div class="download-buttons">
            <a href="/download?type=word" class="btn-word">下載 Word</a>
            <a href="/download?type=pdf" class="btn-pdf">下載 PDF</a>

        </div>

        <h1 class="h1">履歷分析結果</h1>

        <div class="section">
            <div class="with-chart-line">
                <div class="text-content">
    {cleaned_result.replace("三、履歷評分與推薦職缺", f'''三、履歷評分與推薦職缺
    各項評分如下（每項滿分 25 分）：

    - 技能匹配度：{score_data.get("技能匹配度", 0)} 分
    - 工作經驗與成就：{score_data.get("工作經驗與成就", 0)} 分
    - 履歷結構與可讀性：{score_data.get("履歷結構與可讀性", 0)} 分
    - 個人特色：{score_data.get("個人特色", 0)} 分
    - 總分：{score_data.get("總分", 0)} 分（滿分 100 分)''')}
                </div>
                <div class="right-charts">
                    <div class="chart-container">
                        <div class="chart-title">統整分數</div>
                        <canvas id="radarChart"></canvas>
                    </div>
                    <div class="chart-container">
                        <div class="chart-title">DeepSeek 分析</div>
                        <canvas id="chartDeepSeek"></canvas>
                    </div>
                    <div class="chart-container">
                        <div class="chart-title">Gemini 分析</div>
                        <canvas id="chartGemini"></canvas>
                    </div>
                    <div class="chart-container">
                        <div class="chart-title">OpenAI 分析</div>
                        <canvas id="chartOpenAI"></canvas>
                    </div>
                </div>
            </div>
        </div>

        <script>
        function createRadarChart(canvasId, label, scores) {{
            new Chart(document.getElementById(canvasId), {{
                type: 'radar',
                data: {{
                    labels: [
                        '技能匹配度（' + scores[0] + '分）',
                        '工作經驗與成就（' + scores[1] + '分）',
                        '履歷結構與可讀性（' + scores[2] + '分）',
                        '個人特色（' + scores[3] + '分）'
                    ],
                    datasets: [{{
                        label: label,
                        data: scores,
                        fill: true,
                        borderColor: 'rgba(26, 75, 140, 1)',
                        backgroundColor: 'rgba(26, 75, 140, 0.2)',
                        pointBackgroundColor: 'rgba(26, 75, 140, 1)'
                    }}]
                }},
                options: {{
                    responsive: true,
                    plugins: {{
                        legend: {{
                            display: false
                        }}
                    }},
                    scales: {{
                        r: {{
                            min: 0,
                            max: 25,
                            ticks: {{
                                stepSize: 5
                            }}
                        }}
                    }}
                }}
            }});
        }}

        createRadarChart("radarChart", "統整分數", [
            {score_data.get("技能匹配度", 0)},
            {score_data.get("工作經驗與成就", 0)},
            {score_data.get("履歷結構與可讀性", 0)},
            {score_data.get("個人特色", 0)}
        ]);

        createRadarChart("chartDeepSeek", "DeepSeek 分析", [
            {score_data_d.get("技能匹配度", 0)},
            {score_data_d.get("工作經驗與成就", 0)},
            {score_data_d.get("履歷結構與可讀性", 0)},
            {score_data_d.get("個人特色", 0)}
        ]);

        createRadarChart("chartGemini", "Gemini 分析", [
            {score_data_g.get("技能匹配度", 0)},
            {score_data_g.get("工作經驗與成就", 0)},
            {score_data_g.get("履歷結構與可讀性", 0)},
            {score_data_g.get("個人特色", 0)}
        ]);

        createRadarChart("chartOpenAI", "OpenAI 分析", [
            {score_data_o.get("技能匹配度", 0)},
            {score_data_o.get("工作經驗與成就", 0)},
            {score_data_o.get("履歷結構與可讀性", 0)},
            {score_data_o.get("個人特色", 0)}
        ]);
        </script>
    </body>
    </html>
    ''')




@app.route("/download")
def download_result():
    format_type = request.args.get("type", "word")

    try:
        with open(os.path.join(OUTPUT_FOLDER, "result.txt"), "r", encoding="utf-8") as f:
            result_text = f.read()
    except:
        return "找不到分析結果，請先執行分析"

    if not result_text.strip():
        return "無有效分析結果，請返回首頁重新分析"

    if format_type == "pdf":
        file_path = save_analysis_to_pdf(result_text)
        download_name = "履歷分析報告.pdf"
    else:
        file_path = save_analysis_to_word(result_text)
        download_name = "履歷分析報告.docx"

    if not os.path.exists(file_path):
        return "錯誤：檔案建立失敗，無法下載"
    
    print("🚀 準備傳送檔案：", file_path)
    print("📂 檔案是否存在？", os.path.exists(file_path))

    return send_file(
        file_path,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/octet-stream"
    )







if __name__ == "__main__":
    app.run(debug=True)


